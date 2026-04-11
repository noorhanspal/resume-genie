from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import Response
from models.resume import ResumeData, ATSRequest
from services.llm_service import (
    enhance_resume_with_llm, 
    analyze_ats_match, 
    analyze_uploaded_resume,
    enhance_resume_from_text
)
from services.pdf_service import generate_pdf
import io
import pypdf
import docx


router = APIRouter(prefix="/api/resume", tags=["resume"])


@router.post("/enhance")
async def enhance_resume(data: ResumeData):
    try:
        enhanced = enhance_resume_with_llm(data)
        return {"success": True, "data": enhanced}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-pdf")
async def generate_resume_pdf(data: ResumeData):
    try:
        enhanced = enhance_resume_with_llm(data)
        pdf_bytes = generate_pdf(data, enhanced)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{data.personal_info.full_name.replace(" ", "_")}_Resume.pdf"'
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ats-match")
async def ats_match(request: ATSRequest):
    try:
        result = analyze_ats_match(request.resume_data, request.job_description)
        return {"success": True, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/smart-analyze")
async def smart_analyze(file: UploadFile = File(...)):
    try:
        content = await file.read()
        text = ""
        filename = file.filename.lower()
        
        if filename.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
            
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract any text from the file.")
            
        result = analyze_uploaded_resume(text)
        return {"success": True, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload-enhance")
async def upload_enhance(file: UploadFile = File(...), prompt: str = Form(...)):
    try:
        content = await file.read()
        text = ""
        filename = file.filename.lower()
        
        if filename.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
            
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract any text from the file.")
            
        result_data = enhance_resume_from_text(text, prompt)
        return {"success": True, "data": result_data}
    except Exception as e:
        print(f"Error in upload_enhance: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
